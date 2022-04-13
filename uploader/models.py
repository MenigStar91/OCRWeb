from asyncio.windows_events import NULL
from django.db import models
#from core import utils
import hashlib
from PIL import Image
import pytesseract
import docx
import glob,os
import unicodedata
from fpdf import FPDF
# import PyPDF2
import fitz
import random 
import string


def random_value_generator(size=10, chars=string.ascii_uppercase+string.digits):
    return ''.join(random.choice(chars) for _ in range(size))


class FileManager(models.Manager):

    def search(self, query):
        return self.get_queryset().filter(models.Q(internal_reference__icontains=query) |
                                          models.Q(name__icontains=query) |
                                          models.Q(description__icontains=query)
                                          )



class File(models.Model):

    name = models.CharField("Name", max_length=100)
    internal_reference = models.CharField("Internal Reference", max_length=100, editable=False)
    description = models.TextField("Description", blank=True, null=True)
    file = models.FileField(upload_to="OCR_file/input/", verbose_name="Input File")
    create_at = models.DateTimeField("Create at", auto_now_add=True)
    updated_at = models.DateTimeField("Update at", auto_now=True)

    def __str__(self):
        return "{0:03d} - {1}".format(self.id, self.file)


    def text_extraction(self,file):
        img = Image.open(file)
        txt = pytesseract.image_to_string(img, lang='eng')


        my_file = open('static/text.txt', "w+")
        s=''
        for i in txt:
            if(32<=ord(i)<=126 or ord(i)==10):
                s+=i


        my_file.write(s)
        my_file.close()
        
        return my_file

    def txt_to_doc(self):

        tfile=open('static/text.txt','r')
        s=tfile.read()
        tfile.close()
    
        doc=docx.Document()
        doc.add_paragraph(s)
        doc.save("static/text.docx")

        return doc

    def txt_to_pdf(self):
        
        tfile=open('static/text.txt','r')
        s=tfile.read()
        tfile.close()
    
        pdf=FPDF('P','mm',(203,237))
        pdf.add_page()
        pdf.set_font('Arial','', 11)
        pdf.multi_cell(w=0,h=10,txt=s)
        pdf.output(name='static/text.pdf',dest='F')

        return pdf

    def pdf_to_doc(self):
        
        mydoc = docx.Document() # document type

        with fitz.open('media/'+str(self.file)) as doc:
            txt=""
            for page in doc:
                txt+=page.get_text()
        s=''
        for i in txt:
            if(32<=ord(i)<=126 or ord(i)==10):
                s+=i
        doc=docx.Document()
        doc.add_paragraph(txt)
        doc.save('static/text.docx')
        pdf=FPDF('P','mm',(203,237))
        pdf.add_page()
        pdf.set_font('Arial','', 11)
        pdf.multi_cell(w=0,h=10,txt=s)
        pdf.output(name='static/text.pdf',dest='F')

        return mydoc


    def execute_img_and_save_ocr(self):
        
        txtfile=self.text_extraction(self.file)
        docfile=self.txt_to_doc()
        pdffile=self.txt_to_pdf()

        return docfile,pdffile


    def save(self, *args, **kwargs):

        if not self.internal_reference:
            #random_value = utils.random_value_generator(size=20)
            random_value = random_value_generator(size=20)

            while File.objects.filter(internal_reference=random_value).exists():
                #random_value = utils.random_value_generator(size=20)
                random_value = random_value_generator(size=20)
            hash_value = hashlib.md5(bytes(str(self.id) + str(random_value), 'utf-8'))
            self.internal_reference = hash_value.hexdigest()
        super(File, self).save(*args, **kwargs)

    class Meta:
        verbose_name = "PDFFile"
        verbose_name_plural = "PDFFiles"
        ordering = ['id']

    objects = FileManager()